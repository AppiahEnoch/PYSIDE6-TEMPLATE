from common_imports import *
from common_functions import *
logger = logging.getLogger(__name__)
    
def load_json_and_create_document(json_file_path, title):
    """
    Load the JSON file and create the document object with title.
    """
    # Read JSON data
    with open(json_file_path, 'r') as file:
        json_data = json.load(file)

    # Create document
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Add custom styles
    question_style = doc.styles.add_style('Question Style', WD_STYLE_TYPE.PARAGRAPH)
    question_style.font.name = 'Arial'
    question_style.font.size = Pt(12)
    question_style.font.bold = True

    option_style = doc.styles.add_style('Option Style', WD_STYLE_TYPE.PARAGRAPH)
    option_style.font.name = 'Arial'
    option_style.font.size = Pt(11)
    option_style.paragraph_format.left_indent = Inches(0.25)
    option_style.paragraph_format.first_line_indent = Inches(-0.25)

    # Add title
    doc.add_heading(title.upper(), 0)

    return doc, json_data

def write_multiple_choice_questions(doc, json_data):
    """
    Read all multiple-choice questions and write them to the document.
    
    Args:
        doc (Document): The docx Document object to write to.
        json_data (list): List of question dictionaries.
    
    Returns:
        tuple: Updated Document object and list of correct answers.
    """
    correct_answers = []
    question_number = get_config_value('question_number')
    

    for item in json_data:
        if item['type'] == 'Multiple-Choice':
            # Add question
            question_para = doc.add_paragraph(style='Question Style')
            question_para.paragraph_format.space_before = Pt(12)
            question_para.paragraph_format.space_after = Pt(6)
            question_para.add_run(f"{question_number}. {item['question_item']}").bold = True

            # Prepare options
            options = [item['correct_answer']] + [d['detractor_item'] for d in item['detractors']]
            random.shuffle(options)

            # Add options
            for j, option in enumerate(options):
                option_letter = chr(65 + j)
                option_para = doc.add_paragraph(style='Option Style')
                option_para.paragraph_format.left_indent = Inches(0.5)
                option_para.paragraph_format.space_before = Pt(3)
                option_para.paragraph_format.space_after = Pt(3)

                option_run = option_para.add_run(f"({option_letter}) {option}")
                option_run.font.size = Pt(11)

                if option == item['correct_answer']:
                    correct_answers.append((question_number, f"({option_letter}) {option}"))

            # Add space after each question
            doc.add_paragraph().paragraph_format.space_after = Pt(12)

            question_number += 1
            set_config_value('question_number', (question_number))

    return doc, correct_answers

def write_true_false_choice_questions(doc, json_data):
    """
    Read all multiple-choice questions and write them to the document.
    
    Args:
        doc (Document): The docx Document object to write to.
        json_data (list): List of question dictionaries.
    
    Returns:
        tuple: Updated Document object and list of correct answers.
    """
    correct_answers = []
    question_number = get_config_value('question_number')
    

    for item in json_data:
        if item['type'] == 'True/False':
            # Add question
            question_para = doc.add_paragraph(style='Question Style')
            question_para.paragraph_format.space_before = Pt(12)
            question_para.paragraph_format.space_after = Pt(6)
            question_para.add_run(f"{question_number}. {item['question_item']}").bold = True

            # Prepare options
            options = [item['correct_answer']] + [d['detractor_item'] for d in item['detractors']]
            random.shuffle(options)

            # Add options
            for j, option in enumerate(options):
                option_letter = chr(65 + j)
                option_para = doc.add_paragraph(style='Option Style')
                option_para.paragraph_format.left_indent = Inches(0.5)
                option_para.paragraph_format.space_before = Pt(3)
                option_para.paragraph_format.space_after = Pt(3)

                option_run = option_para.add_run(f"({option_letter}) {option}")
                option_run.font.size = Pt(11)

                if option == item['correct_answer']:
                    correct_answers.append((question_number, f"({option_letter}) {option}"))

            # Add space after each question
            doc.add_paragraph().paragraph_format.space_after = Pt(12)

            question_number += 1
            set_config_value('question_number', (question_number))

    return doc, correct_answers

def write_short_answer_questions(doc, json_file_path, start_number):
    """
    Read all short answer questions from a JSON file and write them to the document.
    
    Args:
        doc (Document): The docx Document object to write to.
        json_file_path (str): Path to the JSON file containing question data.
        start_number (int): The starting number for question numbering.
    
    Returns:
        tuple: Updated Document object, list of correct answers, and the last question number.
    """
    try:
        with open(json_file_path, 'r') as file:
            json_data = json.load(file)
    except FileNotFoundError:
        raise FileNotFoundError(f"JSON file not found: {json_file_path}")
    except json.JSONDecodeError:
        raise ValueError(f"Invalid JSON in file: {json_file_path}")

    try:
        short_answer_style = doc.styles.add_style('Short Answer Style', WD_STYLE_TYPE.PARAGRAPH)
        short_answer_style.font.name = 'Arial'
        short_answer_style.font.size = Pt(11)
        short_answer_style.paragraph_format.space_after = Pt(12)
    except ValueError:
        # Style might already exist, so we'll just use it
        short_answer_style = doc.styles['Short Answer Style']

    correct_answers = []
    question_number = get_config_value("question_number")

    for item in json_data:
        if item['type'] == 'Short Answer':
            try:
                question_text = item['question_item']
                question_para = doc.add_paragraph(style='Question Style')
                question_para.add_run(f"{question_number}. {question_text}").bold = True

                for _ in range(3):
                    line_para = doc.add_paragraph(style='Short Answer Style')
                    line_para.add_run('_' * 60)

                doc.add_paragraph().paragraph_format.space_after = Pt(12)

                correct_answer = item['correct_answer']
                correct_answers.append((question_number, correct_answer))
                
                question_number += 1
                
            
            except KeyError as e:
                raise KeyError(f"Missing key in question data: {e}")
            
            set_config_value('question_number', (question_number))

    return doc, correct_answers


import json
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE



def write_essay_questions(doc, json_file_path, start_number):
    start_number = 1
    """
    Read all essay questions from a JSON file and write them to the document.
    
    Args:
        doc (Document): The docx Document object to write to.
        json_file_path (str): Path to the JSON file containing question data.
        start_number (int): The starting number for question numbering.
    
    Returns:
        tuple: Updated Document object, list of correct answers, and the last question number.
    """
    try:
        with open(json_file_path, 'r') as file:
            json_data = json.load(file)
    except FileNotFoundError:
        raise FileNotFoundError(f"JSON file not found: {json_file_path}")
    except json.JSONDecodeError:
        raise ValueError(f"Invalid JSON in file: {json_file_path}")

    # Check if there's a valid Essay type record in the JSON file
    has_essay_questions = any(item['type'] == 'Essay' for item in json_data)

    if has_essay_questions:
        # Add instructions at the top
        instructions = """Section B – Essay Type Questions
Instructions:
- Answer ALL questions from this section.
- Each question carries 20 Marks.
- Write your answers in the provided space in clear and concise language.
- Support your answers with relevant examples, diagrams, or references where applicable.
- Avoid irrelevant details and stay focused on the question asked.
- Time management is crucial. Allocate your time wisely to each question based on the marks assigned.
- Ensure that your handwriting is legible."""
        
        instruction_para = doc.add_paragraph()
        title_run = instruction_para.add_run("Section B – Essay Type Questions")
        title_run.underline = True
        title_run.bold = True
        instruction_para.add_run("\n" + instructions.split("\n", 1)[1])
        instruction_para.style = doc.styles['Normal']
        doc.add_paragraph()  # Add a blank line after instructions
    
    # Define styles
    def get_or_create_style(style_name, base_style=WD_STYLE_TYPE.PARAGRAPH):
        try:
            return doc.styles.add_style(style_name, base_style)
        except ValueError:
            return doc.styles[style_name]

    question_style = get_or_create_style('Question Style')
    question_style.font.name = 'Arial'
    question_style.font.size = Pt(12)
    question_style.font.bold = False  # Remove bold from questions
    question_style.paragraph_format.space_before = Pt(12)
    question_style.paragraph_format.space_after = Pt(6)
    question_style.paragraph_format.left_indent = Inches(0.25)

    correct_answers = []
    question_number = start_number

    for item in json_data:
        if item['type'] == 'Essay':
            try:
                question_text = item['question_item']
                question_para = doc.add_paragraph(style='Question Style')
                
                # Add bold numbering
                number_run = question_para.add_run(f"{question_number}.")
                number_run.bold = True
                
                # Add the question text
                question_para.add_run(f" {question_text}")

                # Add sub-questions if any
                if 'sub_questions' in item:
                    for letter, sub_question in zip('abcdefghijklmnopqrstuvwxyz', item['sub_questions']):
                        sub_para = doc.add_paragraph(style='Question Style')
                        sub_para.paragraph_format.left_indent = Inches(0.5)  # Indent sub-questions
                        
                        # Add bold letter
                        letter_run = sub_para.add_run(f"({letter})")
                        letter_run.bold = True
                        
                        # Add the sub-question text
                        sub_para.add_run(f" {sub_question}")

                # Add two blank lines after each question
                doc.add_paragraph()
                doc.add_paragraph()

                correct_answer = item.get('correct_answer', 'Essay questions typically do not have predefined correct answers.')
                correct_answers.append((question_number, correct_answer))
                
                question_number += 1
                
            except KeyError as e:
                raise KeyError(f"Missing key in question data: {e}")
            
    set_config_value('question_number', question_number)

    return doc, correct_answers







def add_marking_scheme(doc, correct_answers):
    """
    Add a new page and write the marking scheme for the multiple choice items.
    """
    doc.add_page_break()
    doc.add_heading('Marking Scheme', level=1)

    correct_answers.sort(key=lambda x: x[0])

    for question_num, answer in correct_answers:
        answer_para = doc.add_paragraph(style='Normal')
        words = f"{question_num}. {answer}".split()

        processed_words = []
        i = 0
        while i < len(words):
            if i < len(words) - 1 and words[i+1].startswith('.'):
                processed_words.append(words[i] + '\u00A0' + words[i+1])
                i += 2
            else:
                processed_words.append(words[i])
                i += 1

        answer_para.add_run(' '.join(processed_words))

    return doc

def essay_marking_scheme(doc, correct_answers):
    """
    Add a new page and write the marking scheme for essay questions with improved formatting and error handling.
    """
    doc.add_page_break()
    
    # Create and apply heading style
    heading_style = doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.name = 'Arial'
    heading_style.font.size = Pt(16)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph('Essay Questions: Marking Scheme', style='CustomHeading').alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Sort answers by question number, handling different data types
    def get_question_id(item):
        if isinstance(item, dict):
            return item.get('question_id', 0)
        elif isinstance(item, (list, tuple)) and len(item) > 0:
            return item[0] if isinstance(item[0], (int, str)) else 0
        return 0

    correct_answers.sort(key=get_question_id)
    
    for answer_data in correct_answers:
        # Extract question number and answer, handling different data structures
        if isinstance(answer_data, dict):
            question_num = answer_data.get('question_id', 'Unknown')
            answer = answer_data.get('correct_answer', '')
        elif isinstance(answer_data, (list, tuple)) and len(answer_data) >= 2:
            question_num, answer = answer_data[0], answer_data[1]
        else:
            continue  # Skip invalid data
        
        # Add question number
        question_para = doc.add_paragraph()
        question_run = question_para.add_run(f"Question {question_num}")
        question_run.bold = True
        question_run.font.size = Pt(14)
        question_run.font.color.rgb = RGBColor(0, 102, 204)
        
        # Process and format the answer
        if isinstance(answer, str):
            parts = answer.split('\n\n')
        else:
            parts = [str(answer)]  # Convert non-string answers to string
        
        for part in parts:
            part = str(part).strip()
            if part:
                part_para = doc.add_paragraph()
                if part.startswith(('(a)', '(b)', '(c)')):
                    # Part label
                    part_run = part_para.add_run(part[:3])
                    part_run.bold = True
                    part_run.font.size = Pt(12)
                    part_para.add_run(part[3:])
                elif ':' in part:
                    # Definition or key point
                    key, value = part.split(':', 1)
                    key_run = part_para.add_run(key + ':')
                    key_run.bold = True
                    part_para.add_run(value)
                else:
                    # Regular paragraph
                    part_para.add_run(part)
                
                # Set paragraph formatting
                part_para.paragraph_format.left_indent = Pt(18)
                part_para.paragraph_format.line_spacing = 1.15
                part_para.paragraph_format.space_after = Pt(6)
        
        # Add spacing between questions
        doc.add_paragraph()
    
    return doc

def save_and_open_document(doc, output_filename):
    """
    Save the document to the desired path and auto open the document.
    """
    doc.save(output_filename)
    auto_open_folder()
   
    
    